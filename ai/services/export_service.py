from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import io

FONTS = {
    "Times New Roman": "Times New Roman",
    "Gill Sans":       "Gill Sans MT",
    "Arial":           "Arial",
    "Calibri":         "Calibri",
    "Georgia":         "Georgia",
    "Helvetica":       "Arial",
}

def build_docx(
    lines:            list,
    title:            str       = "Transcript",
    content_type:     str       = "both",   # "both" | "transcript" | "translation"
    font_name:        str       = "Calibri",
    font_size:        int       = 12,
    include_speakers: bool      = True,
    include_timestamps: bool    = True,
) -> bytes:
    """
    Build a Word document from transcript lines.
    content_type:
        "both"        — original text + English translation
        "transcript"  — original text only
        "translation" — English translation only
    Returns document as bytes.
    """
    doc = Document()

    # ── Page margins ───────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ──────────────────────────────────────────────────────
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.name  = FONTS.get(font_name, "Calibri")
    title_run.font.size  = Pt(font_size + 6)
    title_run.font.color.rgb = RGBColor(0x1a, 0x7a, 0x4a)  # TranscribeNG green

    doc.add_paragraph()  # spacer

    # ── Subtitle ───────────────────────────────────────────────────
    label = {
        "both":        "Transcript with English Translation",
        "transcript":  "Original Transcript",
        "translation": "English Translation",
    }.get(content_type, "Transcript")

    sub = doc.add_paragraph(label)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.runs[0]
    sub_run.font.name   = FONTS.get(font_name, "Calibri")
    sub_run.font.size   = Pt(font_size - 1)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # ── Transcript lines ───────────────────────────────────────────
    for line in lines:
        speaker   = line.get("speaker", "")
        text      = line.get("text", "")
        timestamp = line.get("time", "")
        lang      = line.get("lang", "")
        translation = line.get("translation")

        # Speaker + timestamp header
        if include_speakers or include_timestamps:
            header_para = doc.add_paragraph()
            if include_speakers:
                speaker_run = header_para.add_run(speaker)
                speaker_run.font.name   = FONTS.get(font_name, "Calibri")
                speaker_run.font.size   = Pt(font_size - 1)
                speaker_run.font.bold   = True
                speaker_run.font.color.rgb = RGBColor(0x1a, 0x7a, 0x4a)

            if include_speakers and include_timestamps:
                sep_run = header_para.add_run("  ·  ")
                sep_run.font.name  = FONTS.get(font_name, "Calibri")
                sep_run.font.size  = Pt(font_size - 2)
                sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            if include_timestamps:
                time_run = header_para.add_run(timestamp)
                time_run.font.name   = FONTS.get(font_name, "Calibri")
                time_run.font.size   = Pt(font_size - 2)
                time_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            if lang and lang.lower() not in ["english", "en"]:
                lang_run = header_para.add_run(f"  [{lang}]")
                lang_run.font.name   = FONTS.get(font_name, "Calibri")
                lang_run.font.size   = Pt(font_size - 2)
                lang_run.font.italic = True
                lang_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Original transcript text
        if content_type in ["both", "transcript"]:
            text_para = doc.add_paragraph()
            text_para.paragraph_format.left_indent = Inches(0.3)
            text_run = text_para.add_run(text)
            text_run.font.name = FONTS.get(font_name, "Calibri")
            text_run.font.size = Pt(font_size)

        # English translation
        if content_type in ["both", "translation"] and translation:
            trans_para = doc.add_paragraph()
            trans_para.paragraph_format.left_indent = Inches(0.3)
            trans_run = trans_para.add_run(
                f"[EN] {translation}" if content_type == "both" else translation
            )
            trans_run.font.name   = FONTS.get(font_name, "Calibri")
            trans_run.font.size   = Pt(font_size)
            trans_run.font.italic = True
            trans_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # Spacer between lines
        doc.add_paragraph()

    # ── Footer ─────────────────────────────────────────────────────
    footer_para = doc.add_paragraph("Generated by TranscribeNG — Real-time transcription for Nigerian languages")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.name   = FONTS.get(font_name, "Calibri")
    footer_run.font.size   = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── Save to bytes ──────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_pdf(lines: list, title: str = "Transcript", **kwargs) -> bytes:
    """Convert docx to PDF using docx2pdf."""
    from docx2pdf import convert

    docx_bytes = build_docx(lines, title, **kwargs)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        tmp_docx.write(docx_bytes)
        tmp_docx_path = tmp_docx.name

    tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")

    try:
        convert(tmp_docx_path, tmp_pdf_path)
        with open(tmp_pdf_path, "rb") as f:
            return f.read()
    finally:
        os.unlink(tmp_docx_path)
        if os.path.exists(tmp_pdf_path):
            os.unlink(tmp_pdf_path)


def build_txt(lines: list, content_type: str = "both", include_speakers: bool = True, include_timestamps: bool = True, **kwargs) -> bytes:
    """Build plain text export."""
    output = []
    for line in lines:
        header_parts = []
        if include_speakers:   header_parts.append(line.get("speaker", ""))
        if include_timestamps: header_parts.append(line.get("time", ""))
        if header_parts:
            output.append("  ".join(header_parts))

        if content_type in ["both", "transcript"]:
            output.append(line.get("text", ""))
        if content_type in ["both", "translation"] and line.get("translation"):
            prefix = "[EN] " if content_type == "both" else ""
            output.append(f"{prefix}{line['translation']}")
        output.append("")

    return "\n".join(output).encode("utf-8")


def build_srt(lines: list, **kwargs) -> bytes:
    """Build SRT subtitle file."""
    def to_srt_time(time_str: str) -> str:
        parts = time_str.split(":")
        if len(parts) == 2:
            return f"00:{parts[0]}:{parts[1]},000"
        return f"{time_str},000"

    output = []
    for i, line in enumerate(lines):
        start = to_srt_time(line.get("time", "00:00"))
        output.append(str(i + 1))
        output.append(f"{start} --> {start}")
        output.append(line.get("text", ""))
        if line.get("translation"):
            output.append(line["translation"])
        output.append("")

    return "\n".join(output).encode("utf-8")